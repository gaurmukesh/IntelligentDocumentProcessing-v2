"""Generate IDP Interview Prep Word Document — with flow, why, challenges, code snippets, key numbers."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if p.runs:
        r = p.runs[0]
        colors = {
            1: RGBColor(0x1F, 0x49, 0x7D),
            2: RGBColor(0x2E, 0x74, 0xB5),
            3: RGBColor(0x70, 0x30, 0xA0),
        }
        r.font.color.rgb = colors.get(level, RGBColor(0x00, 0x00, 0x00))
    return p


def add_para(doc, text, bold=False, italic=False, size=10, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix + "  ")
        r.bold = True
        r.font.size = Pt(9.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(9.5)
    return p


def add_code(doc, lines):
    """Add a grey code block."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, "F3F3F3")
    p = cell.paragraphs[0]
    r = p.add_run(lines)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    doc.add_paragraph()


def add_callout(doc, label, body, bg="FFF2CC", label_color=None):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.size = Pt(9.5)
    if label_color:
        r1.font.color.rgb = label_color
    r2 = p.add_run(body)
    r2.font.size = Pt(9.5)
    doc.add_paragraph()


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        set_cell_bg(c, "1F497D")
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].runs[0].font.size = Pt(9)
    for ri, row in enumerate(rows):
        bg = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            set_cell_bg(c, bg)
            for para in c.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def add_phase_header(doc, project_phase, runtime_label):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    c1 = t.cell(0, 0)
    set_cell_bg(c1, "1F497D")
    r = c1.paragraphs[0].add_run(project_phase)
    r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.size = Pt(9.5)
    c1.width = Inches(2.8)
    c2 = t.cell(0, 1)
    set_cell_bg(c2, "2E74B5")
    r2 = c2.paragraphs[0].add_run(runtime_label)
    r2.bold = True; r2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r2.font.size = Pt(9.5)
    doc.add_paragraph()


# ── Main document ─────────────────────────────────────────────────────────────

def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

    # ── Cover ──────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    t = doc.add_heading("Intelligent Document Processing (IDP)", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    s = doc.add_paragraph("Interview Preparation Guide")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].font.size = Pt(15); s.runs[0].bold = True
    s.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    tag = doc.add_paragraph(
        "End-to-End Flow  •  Why Each Technology  •  Code Snippets  •  Challenges  •  Key Numbers"
    )
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.runs[0].italic = True; tag.runs[0].font.size = Pt(10)
    tag.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    doc.add_paragraph()

    # ── 1. Problem Statement ───────────────────────────────────────────────────
    add_heading(doc, "1. Problem Statement", 1)
    add_para(doc,
        "Universities manually verify student admission documents — checking if the name on the 10th "
        "marksheet matches the Aadhar card, confirming DOBs, validating percentages, and checking "
        "eligibility criteria. This is slow, error-prone, and does not scale."
    )
    add_para(doc,
        "IDP automates this entirely. A student submits three documents and the system extracts all "
        "fields using AI, cross-validates them, checks eligibility rules, and delivers an "
        "APPROVED / REJECTED / REVIEW_REQUIRED decision — no manual intervention required."
    )

    # ── 2. Project Phases ──────────────────────────────────────────────────────
    add_heading(doc, "2. Project Phases — How I Built It (9 Days)", 1)
    add_table(doc,
        ["Phase", "Phase Name", "What Was Built"],
        [
            ("Phase 0", "Project Setup",                        "Monorepo structure, FastAPI venv, Spring Boot via Initializr, SQLite + H2 DB initialisation"),
            ("Phase 1", "Spring Boot Foundation",               "Student & Application entities, JPA repositories, REST controllers (CRUD), DTOs, global exception handler"),
            ("Phase 2", "FastAPI Foundation",                   "FastAPI app, SQLAlchemy async models, document upload endpoint, Pydantic schemas, health check"),
            ("Phase 3", "Document Extraction Service",          "GPT-4o Vision extractor for 10th marksheet, 12th marksheet, Aadhar — per-field confidence scores"),
            ("Phase 4", "Validation Engine",                    "Required field checks, fuzzy name matching, DOB cross-check, percentage eligibility, decision engine"),
            ("Phase 5", "RAG Knowledge Base",                   "LangChain + Qdrant in-memory vector store, eligibility rule ingestion, RAG query + eligibility service"),
            ("Phase 6", "Verification Workflow + Integration",  "Kafka async pipeline, verify endpoints, Spring Boot WebClient integration, ERP callback"),
            ("Phase 7", "Streamlit UI",                         "4 pages: New Application, Applications List, Pipeline Status (live polling), Verification Report"),
            ("Phase 8", "Testing and Stabilisation",            "pytest (FastAPI), JUnit 5 + Mockito (Spring Boot), bug fixes, prompt tuning, edge cases"),
        ],
        col_widths=[0.75, 2.2, 4.0]
    )

    # ── 3. System Architecture ─────────────────────────────────────────────────
    add_heading(doc, "3. System Architecture", 1)
    add_table(doc,
        ["Component", "Technology", "Port", "Role"],
        [
            ("Streamlit UI",         "Python · Streamlit",                   "8501", "Admin frontend — upload, trigger, reports"),
            ("FastAPI AI Service",   "Python · FastAPI · SQLAlchemy async",  "8000", "Document storage, AI extraction, validation"),
            ("Spring Boot ERP",      "Java · Spring Boot · WebFlux",         "8080", "Student/application lifecycle, decision storage"),
            ("Kafka Broker",         "Apache Kafka (KRaft, no Zookeeper)",   "9092", "Async queue — decouples upload from extraction"),
            ("Kafka Consumer",       "Python · kafka-python (separate process)", "—", "Reads messages, calls GPT-4o, saves results"),
            ("GPT-4o Vision",        "OpenAI API",                           "—",    "OCR + structured data extraction from documents"),
            ("Qdrant Vector Store",  "LangChain + Qdrant in-memory",         "—",    "RAG knowledge base for eligibility rules"),
            ("SQLite",               "aiosqlite · SQLAlchemy async",         "—",    "FastAPI DB — documents, extractions, validations"),
            ("H2 Database",          "In-memory JDBC (Spring Boot)",         "—",    "ERP DB — students, applications, decisions"),
        ],
        col_widths=[1.5, 2.0, 0.55, 2.9]
    )

    # ── 4. End-to-End Flow ─────────────────────────────────────────────────────
    add_heading(doc, "4. End-to-End Flow — Step by Step", 1)

    # ── Step 1 ──
    add_phase_header(doc,
        "Phase 1 (Spring Boot Foundation) + Phase 7 (Streamlit UI)",
        "STEP 1 — Create Student & Application")
    add_para(doc, "Admin opens Streamlit, fills in student details and creates an application. The application_id returned is the shared key used across ALL services for this entire flow.")
    add_code(doc,
        "# frontend/utils/api.py\n"
        "def create_student(payload: dict) -> dict:\n"
        "    r = requests.post('http://localhost:8080/api/students', json=payload)\n"
        "    r.raise_for_status()\n"
        "    return r.json()   # returns { 'id': 'student-uuid', 'name': '...' }\n\n"
        "def create_application(student_id: str) -> dict:\n"
        "    r = requests.post('http://localhost:8080/api/applications',\n"
        "                      json={'studentId': student_id})\n"
        "    r.raise_for_status()\n"
        "    return r.json()   # returns { 'id': 'app-uuid', 'status': 'PENDING' }"
    )
    add_callout(doc,
        "Why Spring Boot for the ERP?",
        "Java Spring Boot is the industry standard for enterprise backends. It has mature ORM (JPA), "
        "validation, and transaction management. The ERP would connect to real university systems in "
        "production — Java is the right fit. Python is kept for the AI service where the ML ecosystem lives.",
        bg="DEEAF1", label_color=RGBColor(0x1F, 0x49, 0x7D)
    )

    # ── Step 2 ──
    add_phase_header(doc,
        "Phase 2 (FastAPI Foundation) + Phase 7 (Streamlit UI)",
        "STEP 2 — Upload Documents")
    add_para(doc, "Admin uploads 10th marksheet, 12th marksheet, and Aadhar card. Each upload is independent. FastAPI saves the file, creates a DB record, and immediately drops a Kafka message — the API returns instantly without waiting for extraction.")
    add_code(doc,
        "# backend-fastapi/app/routers/documents.py\n"
        "@router.post('/upload')\n"
        "async def upload_document(file: UploadFile, application_id: str = Form(...),\n"
        "                          doc_type: str = Form(...), db: AsyncSession = Depends(get_db)):\n"
        "    # 1. Save file to disk\n"
        "    file_path = f'uploads/{application_id}/{doc_type}{ext}'\n"
        "    with open(file_path, 'wb') as f:\n"
        "        f.write(await file.read())\n\n"
        "    # 2. Create Document row in SQLite (status = PENDING)\n"
        "    doc = Document(id=str(uuid4()), application_id=application_id,\n"
        "                   doc_type=doc_type, file_path=file_path,\n"
        "                   status=DocumentStatus.PENDING)\n"
        "    db.add(doc); await db.commit()\n\n"
        "    # 3. Publish Kafka message — extraction happens async\n"
        "    producer.send('document-extraction', value={\n"
        "        'document_id': doc.id, 'application_id': application_id,\n"
        "        'doc_type': doc_type, 'file_path': file_path\n"
        "    })\n"
        "    return {'id': doc.id, 'status': 'PENDING'}"
    )
    add_callout(doc,
        "Why FastAPI for the AI service?",
        "FastAPI is async-native (Starlette + asyncio), has automatic OpenAPI docs, and Pydantic "
        "validation built in. It's ideal for a service that makes many async I/O calls (SQLite, OpenAI, "
        "Kafka). Python also has the best AI/ML library ecosystem — langchain, openai, pymupdf are "
        "all Python-first.",
        bg="DEEAF1", label_color=RGBColor(0x1F, 0x49, 0x7D)
    )

    # ── Step 3 ──
    add_phase_header(doc,
        "Phase 3 (Document Extraction Service) + Phase 6 (Verification Workflow)",
        "STEP 3 — AI Extraction via Kafka Consumer")
    add_para(doc, "The Kafka consumer is a separate Python process. It polls for messages, converts PDFs to images, and calls GPT-4o Vision to extract structured data from each document.")
    add_code(doc,
        "# backend-fastapi/app/services/kafka_consumer.py\n"
        "async def process_message(message: dict):\n"
        "    doc.status = DocumentStatus.EXTRACTING\n"
        "    await db.commit()\n\n"
        "    extracted_data, confidence = extract_document(file_path, doc_type)\n"
        "    # → saves ExtractionResult, sets status = EXTRACTED\n\n"
        "# Offset committed ONLY after success — at-least-once delivery\n"
        "asyncio.run(process_message(msg.value))\n"
        "consumer.commit()"
    )
    add_code(doc,
        "# backend-fastapi/app/services/extractor.py\n"
        "def _pdf_first_page_to_base64(file_path):\n"
        "    import fitz  # PyMuPDF — no system dependencies\n"
        "    doc = fitz.open(file_path)\n"
        "    page = doc[0]\n"
        "    mat = fitz.Matrix(3, 3)   # 3x zoom = 216 DPI — sharper digits\n"
        "    pix = page.get_pixmap(matrix=mat)\n"
        "    png_bytes = pix.tobytes('png')\n"
        "    return base64.b64encode(png_bytes).decode('utf-8'), 'image/png'\n\n"
        "def _call_vision(system_prompt, user_prompt, image_b64, media_type):\n"
        "    response = client.chat.completions.create(\n"
        "        model='gpt-4o',\n"
        "        messages=[\n"
        "            {'role': 'system', 'content': system_prompt},\n"
        "            {'role': 'user', 'content': [\n"
        "                {'type': 'text', 'text': user_prompt},\n"
        "                {'type': 'image_url',\n"
        "                 'image_url': {'url': f'data:{media_type};base64,{image_b64}',\n"
        "                               'detail': 'high'}}\n"
        "            ]}\n"
        "        ],\n"
        "        response_format={'type': 'json_object'},\n"
        "        temperature=0, max_tokens=2500\n"
        "    )\n"
        "    return json.loads(response.choices[0].message.content)"
    )
    add_callout(doc,
        "Why Kafka for async processing?",
        "GPT-4o extraction takes 3–5 seconds per document. With 3 documents, synchronous processing "
        "would block the upload API for 15+ seconds. Kafka decouples upload (instant response) from "
        "extraction (background). It also gives at-least-once delivery — offset is committed only "
        "after successful processing. If the consumer crashes mid-extraction, the message is reprocessed "
        "on restart. No documents are ever lost.",
        bg="FFF2CC", label_color=RGBColor(0x7F, 0x60, 0x00)
    )
    add_callout(doc,
        "Why PyMuPDF over pdf2image / poppler?",
        "pdf2image requires poppler — a system C library (brew install poppler / apt install). "
        "This breaks in Docker and CI without extra setup steps. PyMuPDF is a pure Python wheel "
        "with bundled C extensions — zero system dependencies. We also render at 3× zoom (216 DPI) "
        "instead of 2× (144 DPI) because higher resolution lets GPT-4o distinguish similar digits "
        "like '01' vs '03' in month fields.",
        bg="FFF2CC", label_color=RGBColor(0x7F, 0x60, 0x00)
    )
    add_callout(doc,
        "Why GPT-4o Vision over Tesseract or a custom OCR model?",
        "Indian marksheets have wildly inconsistent layouts across CBSE, ICSE, and 20+ state boards. "
        "Training a custom model would need hundreds of labelled samples. GPT-4o Vision understands "
        "document structure without training. It returns structured JSON with per-field confidence "
        "scores in a single API call. Tesseract returns raw text — you'd still need to parse and "
        "map it to fields, and it struggles with mixed-font Indian marksheets.",
        bg="FFF2CC", label_color=RGBColor(0x7F, 0x60, 0x00)
    )

    # ── Step 4 ──
    add_phase_header(doc,
        "Phase 6 (Verification Workflow + Integration)",
        "STEP 4 — Trigger Verification")
    add_para(doc, "After all documents are EXTRACTED, admin clicks Trigger Verification. This flows through Spring Boot first — it must set status to UNDER_REVIEW before delegating to FastAPI.")
    add_code(doc,
        "// backend-springboot — VerificationService.java\n"
        "public void triggerVerification(String applicationId) {\n"
        "    Application app = applicationRepository.findById(applicationId).orElseThrow();\n"
        "    app.setStatus(ApplicationStatus.UNDER_REVIEW);  // ERP owns lifecycle\n"
        "    applicationRepository.save(app);\n\n"
        "    boolean accepted = aiServiceClient.triggerVerification(applicationId);\n"
        "    // Non-blocking WebClient call to FastAPI POST /verify/{applicationId}\n"
        "}"
    )
    add_callout(doc,
        "Why does the trigger go through Spring Boot first?",
        "Spring Boot is the ERP — it owns the application lifecycle. The status must be set to "
        "UNDER_REVIEW in the ERP before we delegate to the AI service. This way, if FastAPI is "
        "temporarily down, the ERP still reflects the correct state. It also means we can swap "
        "out the AI service without changing the ERP's understanding of application statuses.",
        bg="DEEAF1", label_color=RGBColor(0x1F, 0x49, 0x7D)
    )

    # ── Step 5 ──
    add_phase_header(doc,
        "Phase 4 (Validation Engine) + Phase 5 (RAG Knowledge Base)",
        "STEP 5 — Run Validation + RAG Eligibility Check")
    add_para(doc, "A FastAPI background task reads all ExtractionResults and runs 22+ checks across four categories, then runs a RAG eligibility check against the knowledge base.")
    add_code(doc,
        "# backend-fastapi/app/services/validator.py (key checks)\n\n"
        "# A. Required field check — for each doc type\n"
        "def validate_required_fields(extractions, doc_confidences):\n"
        "    for doc_type, required_fields in REQUIRED_FIELDS.items():\n"
        "        for field in required_fields:\n"
        "            value = _val(data.get(field, {}))\n"
        "            status = 'PASS' if value else 'FAIL'\n\n"
        "# B. Fuzzy name match across all document pairs\n"
        "def validate_name_match(extractions):\n"
        "    score = fuzz.token_sort_ratio(name_a, name_b)  # thefuzz library\n"
        "    status = 'PASS' if score >= 85 else ('WARNING' if score >= 70 else 'FAIL')\n\n"
        "# C. DOB cross-check — normalise to DD/MM/YYYY then exact match\n"
        "def validate_dob_match(extractions):\n"
        "    unique_dobs = set(dobs.values())\n"
        "    all_match = len(unique_dobs) == 1  # exact match required\n\n"
        "# D. Percentage eligibility\n"
        "# 10th >= 35%, 12th >= 45%"
    )
    add_code(doc,
        "# backend-fastapi/app/services/rag.py\n"
        "def check_eligibility(course, pct_10th, pct_12th, stream, result_12th):\n"
        "    query = f'Student: stream={stream}, 10th={pct_10th}%, 12th={pct_12th}%, result={result_12th}'\n\n"
        "    # 1. Embed query using text-embedding-3-small\n"
        "    # 2. Qdrant similarity search → top-k relevant eligibility rules\n"
        "    docs = vectorstore.similarity_search(query, k=5)\n\n"
        "    # 3. GPT-4 reasons over retrieved rules\n"
        "    response = llm.invoke(prompt_with_rules_and_profile)\n"
        "    return { 'eligible': True/False, 'reason': '...', 'confidence': 0.0-1.0 }"
    )
    add_callout(doc,
        "Why RAG instead of hardcoded eligibility rules?",
        "Eligibility rules change by course, stream, and year. Hardcoding means a code deploy every "
        "time rules change. With RAG, a non-technical admin updates a text file and re-ingests it — "
        "no code change needed. The LLM also handles nuanced cases like borderline percentages or "
        "stream-specific exceptions that rigid if/else cannot. RAG also provides explainability: "
        "the model reasons over retrieved rules and must give a reason for its decision.",
        bg="FFF2CC", label_color=RGBColor(0x7F, 0x60, 0x00)
    )

    # ── Step 6 ──
    add_phase_header(doc,
        "Phase 6 (Verification Workflow + Integration)",
        "STEP 6 — Save Result & Notify Spring Boot")
    add_para(doc, "After all checks complete, the ValidationResult is saved to SQLite and FastAPI calls back to Spring Boot with the final decision.")
    add_code(doc,
        "# backend-fastapi/app/routers/verify.py\n"
        "async def _notify_erp(application_id, validation_output):\n"
        "    url = f'{settings.erp_base_url}/api/applications/{application_id}/verification-result'\n"
        "    payload = {\n"
        "        'decision':        validation_output['decision'],      # APPROVED / REJECTED\n"
        "        'overall_score':   validation_output['overall_score'], # e.g. 0.91\n"
        "        'decision_reason': validation_output['decision_reason']\n"
        "    }\n"
        "    async with httpx.AsyncClient(timeout=10.0) as client:\n"
        "        response = await client.post(url, json=payload)\n"
        "    # best-effort — failure is logged, not raised"
    )
    add_code(doc,
        "// backend-springboot — VerificationService.java\n"
        "public void applyVerificationResult(String applicationId, VerificationCallbackRequest cb) {\n"
        "    app.setVerificationDecision(mapDecision(cb.getDecision()));\n"
        "    app.setVerificationScore(cb.getOverallScore());\n"
        "    app.setStatus(switch (decision) {\n"
        "        case APPROVED      -> ApplicationStatus.COMPLETED;\n"
        "        case REJECTED      -> ApplicationStatus.REJECTED;\n"
        "        case MANUAL_REVIEW -> ApplicationStatus.UNDER_REVIEW;\n"
        "    });\n"
        "    applicationRepository.save(app);\n"
        "}"
    )
    add_callout(doc,
        "Why a push callback instead of Spring Boot polling?",
        "Polling would mean Spring Boot repeatedly asks FastAPI 'is it done yet?' — wasted requests. "
        "A push callback means Spring Boot is contacted exactly once, when the result is ready. "
        "The callback is best-effort (try/except) because the result is already saved in SQLite. "
        "If Spring Boot is down, the result is not lost — it can be re-triggered manually.",
        bg="DEEAF1", label_color=RGBColor(0x1F, 0x49, 0x7D)
    )

    # ── Step 7 ──
    add_phase_header(doc,
        "Phase 7 (Streamlit UI)",
        "STEP 7 — Monitor Progress & View Report")
    add_para(doc, "Streamlit polls the pipeline-status endpoint every few seconds to show live progress. Once COMPLETE, admin views the full verification report.")
    add_code(doc,
        "# backend-fastapi/app/routers/applications.py\n"
        "# Pipeline stage is DERIVED from DB state — no extra status field needed\n"
        "if DocumentStatus.FAILED in statuses:        pipeline_stage = 'FAILED'\n"
        "elif any(d.status == EXTRACTING ...):        pipeline_stage = 'EXTRACTING'\n"
        "elif any(d.status == PENDING ...):           pipeline_stage = 'UPLOADING'\n"
        "elif validation:                             pipeline_stage = 'COMPLETE'\n"
        "else:                                        pipeline_stage = 'VALIDATING'"
    )
    add_callout(doc,
        "Why derive pipeline_stage from DB state?",
        "Storing pipeline_stage as a separate field creates duplication — it would go out of sync "
        "if document statuses are updated but pipeline_stage is not. Deriving it from the actual "
        "document statuses and presence of a ValidationResult means it is always consistent with "
        "ground truth. One source of truth.",
        bg="E2EFDA", label_color=RGBColor(0x37, 0x5A, 0x23)
    )

    # ── 5. Challenges ──────────────────────────────────────────────────────────
    add_heading(doc, "5. Challenges Faced & How I Solved Them", 1)
    add_para(doc, "These are real problems hit during development — mention them to show the project is genuine.", italic=True)
    doc.add_paragraph()

    challenges = [
        (
            "OpenAI 'Invalid MIME type' error when uploading PDFs",
            "The original code sent the PDF file directly to OpenAI as application/pdf. "
            "OpenAI Vision API only accepts image types (PNG, JPEG). Fix: replaced the broken "
            "fallback with PyMuPDF — the PDF is rendered to a PNG image first, then sent as image/png.",
            "# Before (broken)\n"
            "return base64.b64encode(pdf_bytes), 'application/pdf'  # OpenAI rejects this\n\n"
            "# After (fixed)\n"
            "doc = fitz.open(file_path)\n"
            "pix = doc[0].get_pixmap(matrix=fitz.Matrix(3, 3))\n"
            "return base64.b64encode(pix.tobytes('png')), 'image/png'"
        ),
        (
            "DOB misread: '20/03/2004' instead of '20/01/2004'",
            "GPT-4o misread month '01' (January) as '03' (March) on a 144 DPI render. "
            "The cross-document DOB check correctly caught this as FAIL (Aadhar had the correct date). "
            "Fix: increased zoom from 2× (144 DPI) to 3× (216 DPI) for sharper digit rendering.",
            "mat = fitz.Matrix(2, 2)  # Before — 144 DPI\n"
            "mat = fitz.Matrix(3, 3)  # After  — 216 DPI, clearer digits"
        ),
        (
            "/knowledge-base/status returning 404",
            "FastAPI matched the parameterised route /{application_id}/status before the fixed "
            "path /knowledge-base/status. Fix: moved all fixed-path routes above parameterised routes.",
            "# Wrong order — /{application_id}/status shadowed /knowledge-base/status\n"
            "@router.get('/{application_id}/status')   # registered first — matches 'knowledge-base'\n"
            "@router.get('/knowledge-base/status')     # never reached\n\n"
            "# Fixed order\n"
            "@router.get('/knowledge-base/status')     # fixed paths first\n"
            "@router.get('/{application_id}/status')   # parameterised routes after"
        ),
        (
            "Spring Boot receiving score=null from FastAPI callback",
            "FastAPI sends overall_score (snake_case) but Spring Boot's DTO had overallScore "
            "(camelCase). Jackson couldn't deserialise without a mapping annotation.",
            "// Before — Jackson could not map 'overall_score' → overallScore\n"
            "private Double overallScore;\n\n"
            "// After — explicit mapping\n"
            "@JsonProperty(\"overall_score\")\n"
            "private Double overallScore;\n"
            "@JsonProperty(\"decision_reason\")\n"
            "private String decisionReason;"
        ),
        (
            "12th marksheet: stream and percentage null",
            "CBSE 12th marksheets don't always print stream or a total percentage. "
            "Fix: updated the system prompt to instruct GPT-4o to infer stream from subject names "
            "(Physics + Chemistry + Maths → Science) and calculate percentage from theory subjects "
            "if the total is not printed.",
            "# Added to SYSTEM_12TH prompt:\n"
            "'For stream: infer from subjects if not printed'\n"
            "'  (Physics/Chemistry/Maths → Science)'\n"
            "'  (Accountancy/Business Studies → Commerce)'\n"
            "'For percentage: use printed value; else calculate from'\n"
            "'  total_marks_obtained / total_max_marks × 100'"
        ),
        (
            "Kafka Docker image not available (bitnami/kafka removed from Docker Hub)",
            "bitnami/kafka:3.7 and bitnami/kafka:latest were removed from Docker Hub. "
            "Fix: switched to apache/kafka:latest with updated KRaft-mode environment variables.",
            "# Before (broken)\n"
            "image: bitnami/kafka:3.7\n\n"
            "# After (fixed)\n"
            "image: apache/kafka:latest\n"
            "environment:\n"
            "  KAFKA_NODE_ID: 1\n"
            "  KAFKA_PROCESS_ROLES: broker,controller\n"
            "  KAFKA_LISTENERS: PLAINTEXT://:9092,CONTROLLER://:9093"
        ),
    ]

    for title_text, body, code in challenges:
        add_bullet(doc, body, bold_prefix=title_text)
        add_code(doc, code)

    # ── 6. Key Numbers ─────────────────────────────────────────────────────────
    add_heading(doc, "6. Key Numbers to Mention in Interview", 1)
    add_table(doc,
        ["Metric", "Value", "Why It Matters"],
        [
            ("Project phases",           "9  (Phase 0–8)",        "Shows structured, planned delivery"),
            ("Development time",         "9 days (with AI)",      "vs 18 days without AI assistance"),
            ("Documents supported",      "3 types",               "10th marksheet, 12th marksheet, Aadhar card"),
            ("Validation checks",        "22+ checks",            "Field checks + cross-doc + eligibility + RAG"),
            ("Extraction confidence",    "Per-field (0.0–1.0)",   "Each field has its own confidence score from GPT-4o"),
            ("PDF render resolution",    "216 DPI  (3× zoom)",    "Reduces digit misread errors in OCR"),
            ("Name match threshold",     "≥ 85%",                 "Levenshtein token sort via thefuzz library"),
            ("Kafka delivery guarantee", "At-least-once",         "Offset committed only after successful processing"),
            ("Aadhar privacy",           "Last 4 digits only",    "Full Aadhar number never stored — privacy by design"),
            ("Test frameworks",          "pytest + JUnit 5",      "Both services have unit + integration tests"),
            ("Async HTTP client",        "httpx (FastAPI side)",  "Async callback to Spring Boot without blocking"),
            ("Spring Boot HTTP client",  "WebFlux WebClient",     "Non-blocking HTTP call from Java to FastAPI"),
        ],
        col_widths=[2.0, 1.5, 3.4]
    )

    # ── 7. One-Liner Journey ───────────────────────────────────────────────────
    add_heading(doc, "7. One Document's Complete Journey (For Quick Recall)", 1)
    add_table(doc,
        ["#", "Event", "System", "State After"],
        [
            ("1",  "Admin uploads PDF",                        "Streamlit → FastAPI",        "Document: PENDING  |  SQLite"),
            ("2",  "File saved, Kafka message published",      "FastAPI → Kafka",            "Message in topic: document-extraction"),
            ("3",  "Consumer picks up message",                "Kafka → Consumer",           "Document: EXTRACTING  |  SQLite"),
            ("4",  "PDF rendered to PNG at 216 DPI",           "Consumer (PyMuPDF)",         "PNG bytes in memory"),
            ("5",  "GPT-4o extracts fields + confidence",      "Consumer → OpenAI API",      "Structured JSON with value + confidence per field"),
            ("6",  "ExtractionResult saved",                   "Consumer → SQLite",          "Document: EXTRACTED  |  SQLite"),
            ("7",  "Kafka offset committed",                   "Consumer → Kafka",           "Message acknowledged — no reprocessing"),
            ("8",  "Admin triggers verification",              "Streamlit → Spring Boot",    "Application: UNDER_REVIEW  |  H2"),
            ("9",  "Spring Boot calls FastAPI",                "Spring Boot → FastAPI",      "Background validation task queued"),
            ("10", "22+ validation checks run",                "FastAPI (background)",       "Checks: PASS / FAIL / WARNING"),
            ("11", "RAG eligibility check",                    "FastAPI → Qdrant → OpenAI",  "rag_eligibility_check appended"),
            ("12", "ValidationResult saved",                   "FastAPI → SQLite",           "Decision: APPROVED / REJECTED"),
            ("13", "Spring Boot callback",                     "FastAPI → Spring Boot",      "Application: COMPLETED / REJECTED  |  H2"),
            ("14", "Admin views full report",                  "Streamlit → FastAPI",        "All checks + extracted data displayed"),
        ],
        col_widths=[0.3, 2.5, 2.0, 2.1]
    )

    doc.save("IDP_Interview_Prep.docx")
    print("✓  IDP_Interview_Prep.docx created successfully")


if __name__ == "__main__":
    build()
