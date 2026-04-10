"""Generate Word documents for IDP Application Flow and Phase Connectivity."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:val"), "clear")
    shading_elm.set(qn("w:color"), "auto")
    shading_elm.set(qn("w:fill"), "F3F3F3")
    p._p.get_or_add_pPr().append(shading_elm)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, "2E74B5")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT 1 — Application Flow
# ─────────────────────────────────────────────────────────────────────────────

def build_flow_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("Intelligent Document Processing (IDP)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("End-to-End Application Flow")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    doc.add_paragraph()

    # ── What This System Does ──
    add_heading(doc, "1. What This System Does", 1)
    doc.add_paragraph(
        "Automates the verification of student admission documents — 10th marksheet, "
        "12th marksheet, and Aadhar card — for an Education ERP (Serosoft). The system "
        "uses AI-powered OCR (GPT-4o Vision), cross-document validation, and RAG-based "
        "eligibility checks to produce a verified APPROVED / REJECTED / REVIEW_REQUIRED decision."
    )

    # ── Architecture ──
    add_heading(doc, "2. Architecture Overview", 1)
    add_table(doc,
        ["Component", "Technology", "Port", "Role"],
        [
            ("Streamlit UI",      "Python · Streamlit",                      "8501", "Admin frontend — upload, trigger, view reports"),
            ("FastAPI AI Service","Python · FastAPI · SQLAlchemy",            "8000", "Document storage, extraction orchestration, validation"),
            ("Spring Boot ERP",   "Java · Spring Boot · JPA · WebFlux",       "8080", "Student & application lifecycle management"),
            ("Kafka Broker",      "Apache Kafka (KRaft — no Zookeeper)",       "9092", "Async message queue for document processing jobs"),
            ("Kafka Consumer",    "Python process (kafka-python)",             "—",    "Reads from Kafka, calls GPT-4o, saves extraction results"),
            ("GPT-4o Vision",     "OpenAI API",                               "—",    "OCR extraction from PDF/image documents"),
            ("Qdrant",            "In-memory vector store (LangChain)",        "—",    "RAG knowledge base for eligibility rules"),
            ("SQLite",            "aiosqlite · SQLAlchemy async",             "—",    "FastAPI persistence (documents, extractions, validations)"),
            ("H2",                "In-memory JDBC (Spring Boot)",             "—",    "ERP persistence (students, applications, decisions)"),
        ],
        col_widths=[1.5, 2.0, 0.6, 2.8]
    )
    doc.add_paragraph()

    # ── Technology Stack ──
    add_heading(doc, "3. Technology Stack", 1)
    add_table(doc,
        ["Layer", "Technology"],
        [
            ("Frontend",       "Python · Streamlit"),
            ("AI Service",     "Python · FastAPI · SQLAlchemy (async) · aiosqlite"),
            ("ERP Backend",    "Java · Spring Boot · Spring Data JPA · WebFlux WebClient"),
            ("Message Queue",  "Apache Kafka (KRaft mode — no Zookeeper)"),
            ("OCR / Extraction","OpenAI GPT-4o Vision API"),
            ("PDF Rendering",  "PyMuPDF (fitz) — converts PDF pages to PNG at 216 DPI"),
            ("RAG",            "LangChain · OpenAI text-embedding-3-small · Qdrant"),
            ("Name Matching",  "thefuzz (Levenshtein distance — token sort ratio)"),
            ("Async HTTP",     "httpx (FastAPI → Spring Boot callback)"),
            ("Spring HTTP",    "WebFlux WebClient (Spring Boot → FastAPI)"),
        ],
        col_widths=[2.0, 5.0]
    )
    doc.add_paragraph()

    # ── End-to-End Flow ──
    add_heading(doc, "4. End-to-End Flow", 1)

    # Phase 1
    add_heading(doc, "Phase 1 — Student & Application Setup", 2)
    doc.add_paragraph("The admin creates a student record and an application in the Spring Boot ERP via the Streamlit UI.")
    add_table(doc,
        ["Step", "Action", "Endpoint", "Service"],
        [
            ("1", "Create Student",      "POST /api/students",      "Spring Boot"),
            ("2", "Create Application",  "POST /api/applications",  "Spring Boot"),
        ],
        col_widths=[0.4, 1.8, 2.4, 1.5]
    )
    doc.add_paragraph()
    doc.add_paragraph("The application_id (UUID) returned here is the shared key used across all subsequent steps.")

    # Phase 2
    add_heading(doc, "Phase 2 — Document Upload", 2)
    doc.add_paragraph("The admin uploads three documents: 10th marksheet, 12th marksheet, and Aadhar card.")
    add_table(doc,
        ["Step", "Action", "Detail"],
        [
            ("1", "Upload file",          "POST /documents/upload  (FastAPI) — multipart with file, application_id, doc_type"),
            ("2", "Save to disk",         "uploads/{application_id}/{doc_type}.pdf"),
            ("3", "Save Document row",    "SQLite — status = PENDING"),
            ("4", "Publish Kafka message","topic: document-extraction  |  payload: document_id, application_id, doc_type, file_path"),
        ],
        col_widths=[0.4, 1.8, 4.7]
    )

    # Phase 3
    add_heading(doc, "Phase 3 — AI Extraction (Async via Kafka)", 2)
    doc.add_paragraph("The Kafka consumer (separate Python process) picks up each message and calls GPT-4o Vision.")
    add_table(doc,
        ["Step", "Action", "Detail"],
        [
            ("1", "Poll Kafka",           "Consumer polls topic every 1 second"),
            ("2", "Status → EXTRACTING", "Document row updated in SQLite"),
            ("3", "PDF → PNG",            "PyMuPDF renders first page at 216 DPI (3× zoom)"),
            ("4", "Base64 encode",        "PNG bytes → base64 string"),
            ("5", "GPT-4o Vision call",   "System prompt + user prompt (JSON schema) + image → structured JSON"),
            ("6", "Save result",          "ExtractionResult row saved in SQLite with confidence score"),
            ("7", "Status → EXTRACTED",  "Document row updated (or FAILED on error)"),
            ("8", "Kafka commit",         "Offset committed only after successful processing"),
        ],
        col_widths=[0.4, 1.8, 4.7]
    )
    doc.add_paragraph()

    add_heading(doc, "Extracted Fields per Document Type", 3)
    add_table(doc,
        ["Document", "Fields Extracted"],
        [
            ("MARKSHEET_10TH", "student_name, date_of_birth, school_name, board, exam_year, roll_number, subjects[], total_marks, percentage, result, grade"),
            ("MARKSHEET_12TH", "student_name, date_of_birth*, board, stream (inferred if not printed), school_name, exam_year, roll_number, subjects[], total_marks (theory only), percentage, result"),
            ("AADHAR",         "full_name, date_of_birth, gender, aadhar_last4 (last 4 digits only), address, is_front_side"),
        ],
        col_widths=[1.6, 5.3]
    )
    doc.add_paragraph("* CBSE 12th marksheets typically do not print date_of_birth.")

    # Phase 4
    add_heading(doc, "Phase 4 — Verification Trigger", 2)
    doc.add_paragraph(
        "After all documents are extracted, the admin triggers verification. "
        "This flows through Spring Boot (which owns the application lifecycle) before reaching FastAPI."
    )
    add_table(doc,
        ["Step", "From", "To", "Endpoint", "Effect"],
        [
            ("1", "Streamlit",    "Spring Boot", "POST /api/applications/{id}/trigger-verification", "Application status → UNDER_REVIEW"),
            ("2", "Spring Boot",  "FastAPI",     "POST /verify/{application_id}",                    "Validation background task queued"),
            ("3", "FastAPI",      "—",           "(background task)",                                "Response returned immediately; validation runs async"),
        ],
        col_widths=[0.4, 1.1, 1.1, 2.8, 2.0]
    )

    # Phase 5
    add_heading(doc, "Phase 5 — Validation Engine", 2)
    doc.add_paragraph("The background task reads all extraction results and runs four categories of checks.")
    add_table(doc,
        ["Check Category", "Checks Run", "Pass Criteria"],
        [
            ("A. Required Fields",      "One check per required field per document",                          "Field value is not null"),
            ("B. Cross-doc Name Match", "All pairs: 10th ↔ 12th, 10th ↔ Aadhar, 12th ↔ Aadhar",            "Fuzzy match ≥ 85% (Levenshtein token sort)"),
            ("C. Cross-doc DOB Match",  "10th DOB ↔ Aadhar DOB (normalised to DD/MM/YYYY)",                  "Exact match after normalisation"),
            ("D. Percentage Eligibility","10th ≥ 35%,  12th ≥ 45%",                                         "Percentage meets threshold"),
            ("E. Result Check",         "12th result field",                                                  "Value = PASS"),
            ("F. RAG Eligibility",      "Student profile queried against Qdrant knowledge base via GPT-4",    "eligible = true"),
        ],
        col_widths=[1.8, 2.8, 2.3]
    )
    doc.add_paragraph()

    add_heading(doc, "Decision Logic", 3)
    add_table(doc,
        ["Decision", "Condition"],
        [
            ("APPROVED",        "All critical checks PASS"),
            ("REVIEW_REQUIRED", "Warnings present but no critical failures"),
            ("REJECTED",        "One or more critical checks FAIL (DOB mismatch, missing field, result=FAIL, name mismatch)"),
        ],
        col_widths=[1.8, 5.1]
    )

    # Phase 6
    add_heading(doc, "Phase 6 — Result Callback & Report", 2)
    add_table(doc,
        ["Step", "Action", "Detail"],
        [
            ("1", "Save ValidationResult", "SQLite — checks[], overall_score, decision, decision_reason"),
            ("2", "Callback to Spring Boot","POST /api/applications/{id}/verification-result  (httpx, best-effort)"),
            ("3", "Spring Boot updates ERP","Maps APPROVED→COMPLETED, REJECTED→REJECTED, MANUAL_REVIEW→UNDER_REVIEW"),
            ("4", "Streamlit polls status", "GET /applications/{id}/pipeline-status  →  pipeline_stage = COMPLETE"),
            ("5", "Admin views report",     "GET /verify/{id}/report  →  full report with all checks and extracted data"),
        ],
        col_widths=[0.4, 2.0, 4.5]
    )

    # ── Document Types Supported ──
    add_heading(doc, "5. Document Types Supported", 1)
    add_table(doc,
        ["Document", "Key Fields", "Cross-Checks Performed"],
        [
            ("10th Marksheet", "Name, DOB, Board, School, Roll No, Subjects, %, Result", "Name ↔ 12th ↔ Aadhar  |  DOB ↔ Aadhar"),
            ("12th Marksheet", "Name, Board, Stream, School, Roll No, Subjects, %, Result", "Name ↔ 10th ↔ Aadhar"),
            ("Aadhar Card",    "Name, DOB, Gender, Last-4 digits, Address", "Name ↔ 10th ↔ 12th  |  DOB ↔ 10th"),
        ],
        col_widths=[1.5, 2.8, 2.6]
    )
    doc.add_paragraph()

    # ── Key Design Decisions ──
    add_heading(doc, "6. Key Design Decisions", 1)
    add_table(doc,
        ["Decision", "Rationale"],
        [
            ("Kafka for async processing",     "Document AI extraction takes 2–5 s; async keeps upload API responsive"),
            ("PyMuPDF for PDF→PNG",            "No system dependencies (unlike pdf2image + poppler); pure Python wheel"),
            ("GPT-4o Vision",                  "Best-in-class OCR for complex Indian marksheet layouts"),
            ("216 DPI rendering (3× zoom)",    "Reduces digit misreads (e.g. '1' vs '3' in month field)"),
            ("Aadhar last-4 only",             "Privacy by design — full Aadhar number never stored"),
            ("Fuzzy name matching",            "Handles case differences, initials, minor spelling variations"),
            ("RAG for eligibility rules",      "Rules in knowledge base can be updated without code changes"),
            ("Spring Boot callback (push)",    "ERP does not poll; FastAPI pushes result when ready — decoupled"),
            ("At-least-once Kafka delivery",   "Offset committed only after successful processing — no lost messages"),
        ],
        col_widths=[2.2, 4.7]
    )
    doc.add_paragraph()

    # ── Running the System ──
    add_heading(doc, "7. Running the System", 1)
    add_table(doc,
        ["Terminal", "Command", "Purpose"],
        [
            ("Terminal 1", "docker compose up -d",                                              "Start Kafka broker"),
            ("Terminal 2", "cd backend-fastapi && source venv/bin/activate && uvicorn app.main:app --reload --port 8000", "FastAPI AI service"),
            ("Terminal 3", "cd backend-fastapi && source venv/bin/activate && python -m app.services.kafka_consumer",     "Kafka consumer process"),
            ("Terminal 4", "cd backend-springboot && ./mvnw spring-boot:run",                   "Spring Boot ERP"),
            ("Terminal 5", "cd frontend && streamlit run Home.py",                              "Streamlit UI"),
        ],
        col_widths=[1.0, 3.8, 2.1]
    )
    doc.add_paragraph()
    doc.add_paragraph("One-time setup after first start:")
    add_code(doc, "curl -X POST http://localhost:8000/verify/knowledge-base/ingest")
    doc.add_paragraph("This loads the eligibility rules into Qdrant for RAG checks.")

    doc.save("IDP_Application_Flow.docx")
    print("✓ IDP_Application_Flow.docx saved")


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT 2 — Phase Connectivity
# ─────────────────────────────────────────────────────────────────────────────

def build_connectivity_doc():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading("IDP — Phase Connectivity & API Call Chain", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("How each phase connects to the next — exact endpoints, data hand-offs, and code entry points")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    doc.add_paragraph()

    # ── Shared Key ──
    add_heading(doc, "1. The Shared Key — application_id", 1)
    doc.add_paragraph(
        "Every system uses the application_id (UUID) generated by Spring Boot as the shared identifier. "
        "It is present in every Kafka message, every SQLite row, every API call, and every HTTP callback."
    )
    add_table(doc,
        ["ID", "Created By", "Used By"],
        [
            ("student_id",      "Spring Boot (student create)",      "Spring Boot — links application to student"),
            ("application_id",  "Spring Boot (application create)",  "ALL services — Streamlit, FastAPI, Kafka, Spring Boot"),
            ("document_id",     "FastAPI (file upload)",             "Kafka message, ExtractionResult, status updates"),
        ],
        col_widths=[1.5, 2.2, 3.2]
    )
    doc.add_paragraph()

    # ── Phase 1 → 2 ──
    add_heading(doc, "2. Phase 1 → Phase 2: Setup enables Upload", 1)
    doc.add_paragraph(
        "After the admin creates a student and application in Spring Boot, the application_id is stored "
        "in Streamlit session state. Every subsequent upload call carries this ID."
    )
    add_table(doc,
        ["Step", "Caller", "Endpoint", "Key Data Passed"],
        [
            ("Create student",      "Streamlit → Spring Boot", "POST /api/students",     "name, email, phone, dateOfBirth"),
            ("Create application",  "Streamlit → Spring Boot", "POST /api/applications", "studentId"),
            ("Receive app_id",      "Spring Boot → Streamlit", "(response body)",        "application_id  ← stored in session"),
        ],
        col_widths=[1.6, 2.0, 2.2, 2.1]
    )
    doc.add_paragraph()

    # ── Phase 2 → 3 ──
    add_heading(doc, "3. Phase 2 → Phase 3: Upload triggers Extraction", 1)
    doc.add_paragraph(
        "Uploading a file creates a Document record and drops a Kafka message. "
        "The Kafka consumer is a separate Python process that runs independently — "
        "it does not share memory with the FastAPI web server."
    )
    add_table(doc,
        ["Step", "Caller", "Endpoint / Target", "Key Data"],
        [
            ("Upload file",          "Streamlit → FastAPI",       "POST /documents/upload",          "file bytes, application_id, doc_type"),
            ("Save to disk",         "FastAPI (internal)",         "uploads/{app_id}/{doc_type}.pdf", "file_path stored in Document row"),
            ("Create Document row",  "FastAPI → SQLite",           "(ORM insert)",                   "id, application_id, doc_type, status=PENDING"),
            ("Publish Kafka message","FastAPI → Kafka broker",     "topic: document-extraction",      "document_id, application_id, doc_type, file_path"),
            ("Poll message",         "Consumer ← Kafka broker",    "poll(timeout_ms=1000)",           "same payload as above"),
            ("PDF → PNG",            "Consumer (PyMuPDF)",         "(in-process)",                   "3× zoom, 216 DPI → base64 PNG"),
            ("GPT-4o call",          "Consumer → OpenAI API",      "POST /v1/chat/completions",       "base64 image + JSON schema prompt"),
            ("Save ExtractionResult","Consumer → SQLite",          "(ORM insert)",                   "doc_type, extracted_data (JSON), confidence_score"),
            ("Kafka commit",         "Consumer → Kafka broker",    "consumer.commit()",               "committed only on success — no data loss"),
        ],
        col_widths=[1.7, 1.8, 2.2, 2.2]
    )
    doc.add_paragraph()

    # ── Phase 3 → 4 ──
    add_heading(doc, "4. Phase 3 → Phase 4: Extraction enables Verification", 1)
    doc.add_paragraph(
        "Once all documents reach EXTRACTED status, the admin triggers verification. "
        "The trigger flows through Spring Boot first because Spring Boot owns the application lifecycle "
        "— it must set status to UNDER_REVIEW before delegating to FastAPI."
    )
    add_table(doc,
        ["Step", "Caller", "Endpoint", "Effect"],
        [
            ("Trigger verification",   "Streamlit → Spring Boot",  "POST /api/applications/{id}/trigger-verification",  "Spring Boot: status → UNDER_REVIEW"),
            ("Delegate to FastAPI",    "Spring Boot → FastAPI",    "POST /verify/{application_id}",                     "Via WebClient (non-blocking WebFlux)"),
            ("Queue background task",  "FastAPI (internal)",       "(background_tasks.add_task)",                       "Returns 202 immediately; validation runs async"),
        ],
        col_widths=[1.8, 1.8, 2.6, 1.7]
    )
    doc.add_paragraph()

    # ── Phase 4 → 5 ──
    add_heading(doc, "5. Phase 4 → Phase 5: Trigger runs Validation", 1)
    doc.add_paragraph(
        "The background task _run_validation_task reads all ExtractionResults, runs checks, "
        "and saves a ValidationResult. This runs entirely within FastAPI's process."
    )
    add_table(doc,
        ["Step", "Source", "Action"],
        [
            ("1", "SQLite",                    "Read ExtractionResult rows for application_id"),
            ("2", "SQLite",                    "Read Document rows — check all are EXTRACTED (not PENDING/EXTRACTING)"),
            ("3", "validator.py",              "validate_required_fields() — one check per field per doc"),
            ("4", "validator.py",              "validate_name_match() — fuzzy Levenshtein across all doc pairs"),
            ("5", "validator.py",              "validate_dob_match() — exact match after DD/MM/YYYY normalisation"),
            ("6", "validator.py",              "validate_percentage_eligibility() — 10th ≥ 35%, 12th ≥ 45%"),
            ("7", "rag.py (if KB populated)",  "check_eligibility() — embed profile, Qdrant search, GPT-4 reasoning"),
            ("8", "SQLite",                    "Save ValidationResult — checks[], overall_score, decision, decision_reason"),
        ],
        col_widths=[0.4, 2.0, 4.5]
    )
    doc.add_paragraph()

    # ── Phase 5 → 6 ──
    add_heading(doc, "6. Phase 5 → Phase 6: Validation notifies Spring Boot", 1)
    doc.add_paragraph(
        "After saving the ValidationResult, FastAPI makes a best-effort HTTP callback to Spring Boot. "
        "If the callback fails (Spring Boot is down), it is logged but does not fail the validation."
    )
    add_table(doc,
        ["Step", "Caller", "Endpoint", "Payload / Effect"],
        [
            ("POST callback",        "FastAPI → Spring Boot",    "POST /api/applications/{id}/verification-result",   "decision, overall_score, decision_reason"),
            ("Map decision → status","Spring Boot (internal)",   "VerificationService.applyVerificationResult()",     "APPROVED→COMPLETED, REJECTED→REJECTED, MANUAL_REVIEW→UNDER_REVIEW"),
            ("Save to H2",           "Spring Boot → H2 DB",      "(JPA save)",                                        "verificationDecision, verificationScore, decisionReason, status updated"),
        ],
        col_widths=[1.5, 1.8, 2.6, 2.0]
    )
    doc.add_paragraph()

    # ── Phase 6 → UI ──
    add_heading(doc, "7. Phase 6 → UI: Streamlit reads the result", 1)
    doc.add_paragraph("Streamlit polls FastAPI every few seconds to show live pipeline progress, then fetches the full report.")
    add_table(doc,
        ["Step", "Caller", "Endpoint", "pipeline_stage returned"],
        [
            ("Poll status", "Streamlit → FastAPI", "GET /applications/{id}/pipeline-status", "UPLOADING / EXTRACTING / VALIDATING / COMPLETE / FAILED"),
            ("Fetch report","Streamlit → FastAPI", "GET /verify/{id}/report",                "Full JSON: decision, score, all checks, extracted data"),
        ],
        col_widths=[0.7, 1.8, 2.5, 2.9]
    )
    doc.add_paragraph()
    doc.add_paragraph("pipeline_stage is derived from database state — no extra status field needed:")
    add_table(doc,
        ["Condition (checked in order)", "pipeline_stage"],
        [
            ("Any Document has status = FAILED",       "FAILED"),
            ("Any Document has status = EXTRACTING",   "EXTRACTING"),
            ("Any Document has status = PENDING",      "UPLOADING"),
            ("ValidationResult row exists",            "COMPLETE"),
            ("Otherwise (all EXTRACTED, no validation yet)", "VALIDATING"),
        ],
        col_widths=[3.8, 1.5]
    )
    doc.add_paragraph()

    # ── Full Journey ──
    add_heading(doc, "8. One Document's Complete Journey", 1)
    add_table(doc,
        ["#", "Event", "State Change", "Where stored"],
        [
            ("1",  "Admin uploads 10th_marksheet.pdf",           "Document: PENDING",              "FastAPI SQLite"),
            ("2",  "Kafka message published",                     "(in Kafka topic)",               "Kafka broker"),
            ("3",  "Consumer picks up message",                   "Document: EXTRACTING",           "FastAPI SQLite"),
            ("4",  "PyMuPDF converts PDF → PNG",                  "(in memory)",                    "—"),
            ("5",  "GPT-4o Vision extracts fields",               "(API call)",                     "OpenAI"),
            ("6",  "ExtractionResult saved",                      "Document: EXTRACTED",            "FastAPI SQLite"),
            ("7",  "Kafka offset committed",                      "(no reprocessing)",              "Kafka broker"),
            ("8",  "Admin triggers verification",                 "Application: UNDER_REVIEW",      "Spring Boot H2"),
            ("9",  "FastAPI validation task runs",                "(background)",                   "—"),
            ("10", "ValidationResult saved (APPROVED)",           "(result persisted)",             "FastAPI SQLite"),
            ("11", "Spring Boot callback received",               "Application: COMPLETED",         "Spring Boot H2"),
            ("12", "Admin views report in Streamlit",             "(read-only)",                    "FastAPI SQLite (read)"),
        ],
        col_widths=[0.3, 2.5, 2.0, 2.1]
    )
    doc.add_paragraph()

    # ── Error Handling ──
    add_heading(doc, "9. Error Handling at Each Boundary", 1)
    add_table(doc,
        ["Boundary", "Error Type", "Handling"],
        [
            ("Streamlit → Spring Boot",     "HTTP 4xx / 5xx",      "raise_for_status() → error shown in UI"),
            ("Streamlit → FastAPI",         "HTTP 4xx / 5xx",      "raise_for_status() → error shown in UI"),
            ("FastAPI → Kafka (produce)",   "KafkaProducerError",   "Exception propagated → 500 returned to Streamlit"),
            ("Kafka Consumer → OpenAI",     "Any exception",        "ExtractionResult saved with error_message; Document status = FAILED; Kafka NOT committed if crash"),
            ("FastAPI → Spring Boot callback", "httpx exception",  "Logged as WARNING — non-blocking, does not fail validation"),
            ("Spring Boot → FastAPI (WebClient)", "Non-2xx / timeout", "Logged as warning — does not block ERP response to Streamlit"),
        ],
        col_widths=[2.0, 1.8, 3.1]
    )

    doc.save("IDP_Phase_Connectivity.docx")
    print("✓ IDP_Phase_Connectivity.docx saved")


if __name__ == "__main__":
    build_flow_doc()
    build_connectivity_doc()
